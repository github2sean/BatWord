# 官方文档：https://python-docx.readthedocs.org/en/latest/
# !/usr/bin/env python
# coding: utf-8
import random
from datetime import datetime

from translate import Translator
from JavaAnnotation import JavaAnnotation
from docx import Document
from docx.shared import Pt
import re


class OptionDocs:
    # 缓存usefulField
    cacheList = []


def handDocs(codePath, descPath):
    # 创建 Document 对象，相当于打开一个 word 文档
    codePath = '/Users/sean/Downloads/房屋雨水收集利用回收管理系统软件/1  java  房屋雨水收集利用回收管理系统软件.docx'
    descPath = '/Users/sean/Downloads/房屋雨水收集利用回收管理系统软件/房屋雨水收集利用回收管理系统软件.docx'
    document = Document(codePath)
    paragraph = document.paragraphs
    docLen = len(paragraph)
    print("paragraphs 长度:" + str(docLen))
    # 以下是将简单句子从英语翻译中文
    translator = Translator("chinese")
    usefulInfo = findUsefulFiled(descPath)
    functionList = usefulInfo['functionList']
    classList = usefulInfo['classList']
    pageLine = 30
    before10Page = 30 * 10
    totalPage = docLen / 30
    after5PAge = docLen - pageLine * 5
    lastClassRow = 0
    lastFunctionRow = 0
    for row, p in enumerate(paragraph):

        for r in p.runs:
            content = r.text
            if row < before10Page or row > after5PAge:

                # 添加类注释
                if (lastClassRow == 0 or row - lastClassRow > 5) and 'class ' in content and ' ' in content:
                    lastClassRow = row
                    addClassAnnotation(p, content, translator, classList)

                # 添加方法注释
                if isFunctionHead(content) and (
                        lastFunctionRow == 0 or row - lastFunctionRow > 5) and '(' in content and ')' in content and isFunctionTail(
                    content):
                    lastFunctionRow = row
                    addFunctionAnnotation(p, content, translator, functionList)
            # 添加固定日志输出注释
            addLogAnnotation(p, content)

            # 添加创建实例注释
            addNewClass(p, content)

            # 提交事务注释
            addCommit(p, content)

            # 回滚事务注释
            addRoolBack(p, content)

    # 如果还有剩余文字则随机注释
    if len(functionList) > 0:
        for r2, p2 in enumerate(paragraph):
            # 随机注释
            if r2 % (random.randint(100, 200)) == 10:
                addSingleLineAnnotation(p2, functionList)
        # print(str(row)+'========')

    document.save("/Users/sean/Downloads/房屋雨水收集利用回收管理系统软件/1  java  房屋雨水收集利用回收管理系统软件_" + getNowTime() + ".docx")
    return


def findUsefulFiled(path):
    functionList = []
    classList = []
    path = "/Users/sean/Downloads/2会计总帐专业版智能管理系统/2会计总帐专业版智能管理系统.docx"
    infoDocument = Document(path)
    for p in infoDocument.paragraphs:
        if p.style.name == 'Heading 2':
            classList.append(p.text)
        elif p.style.name == 'Normal':
            for r in p.runs:
                if True and len(r.text) > 44:
                    functionList.append(r.text)
    print(functionList)
    print(classList)
    result = {'functionList': functionList, 'classList': classList}
    return result


def matchRule():
    return False


def extractClassName(originStr, flag):
    result = None
    if originStr and flag in originStr:
        nameList = originStr.split(flag)
        for row, name in enumerate(nameList):
            if 'class' in name and len(nameList) >= row + 2 and '{' in originStr:
                result = nameList[row + 1].replace('{', '')
            elif 'class' in name:
                originStr = originStr.rstrip()
                result = nameList[row + 1]
    print('extractClassName:', result)
    return result


def extractFunctionName(originStr, flag):
    result = None
    if originStr and flag in originStr:
        nameList = originStr.split(flag)
        print(nameList)
        for row, name in enumerate(nameList):
            if '()' in name:
                return name.replace('()', '')
            elif '(' in name and ')' in name:
                return nameList[row - 1]
    return result


def handDesc(returnList: list, originStr: str):
    if originStr:
        handStrLen = len(originStr)
        print('handStrLen:%d' % handStrLen)
        if handStrLen > 34:
            strHead = originStr[:34]
            strTail = originStr[34:]
            returnList.append(strHead)
            handDesc(returnList, strTail)
        else:
            returnList.append(originStr)
    return returnList


def addClassAnnotation(p, content, translator, classList):
    name = extractClassName(content, ' ')
    params = []
    javaAnnotation = JavaAnnotation(params)
    p.insert_paragraph_before(javaAnnotation.head, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.author, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.version, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.className + name + ' 类', p.style).paragraph_format.left_indent = Pt(20)
    classDesc = classList.pop(0)
    print('classDesc :', classDesc)
    print('classList :', classList)
    p.insert_paragraph_before('* 描述 ' + classDesc, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.tail, p.style).paragraph_format.left_indent = Pt(20)
    print('addClassAnnotation executing....')


def addFunctionAnnotation(p, content, translator, functionList):
    name = extractFunctionName(content, ' ')
    print('functionName:', name)
    params = []
    javaAnnotation = JavaAnnotation(params)
    p.insert_paragraph_before(javaAnnotation.head, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.lineSplit + '[' + name + ']', p.style).paragraph_format.left_indent = Pt(
        20)
    returnList = []
    functionDesc = functionList.pop(0)
    print('functionDesc:' + functionDesc)
    for row, up in enumerate(handDesc(returnList, functionDesc)):
        if row == 0:
            p.insert_paragraph_before('* 描述 ' + up, p.style).paragraph_format.left_indent = Pt(20)
        else:
            p.insert_paragraph_before('* ' + up, p.style).paragraph_format.left_indent = Pt(20)
    p.insert_paragraph_before(javaAnnotation.tail, p.style).paragraph_format.left_indent = Pt(20)
    print('addFunctionAnnotation executing....')


def addSingleLineAnnotation(p, descList):
    if descList:
        lineDesc = descList.pop(0)
        listVolume = []
        for row, up in enumerate(handDesc(listVolume, lineDesc)):
            if row == 0:
                p.insert_paragraph_before('/*  ' + up, p.style).paragraph_format.left_indent = Pt(20)
            else:
                p.insert_paragraph_before('* ' + up, p.style).paragraph_format.left_indent = Pt(20)
        p.insert_paragraph_before('*/ ' + up, p.style).paragraph_format.left_indent = Pt(20)

def addLogAnnotation(p, originStr):
    if 'logger.info(' in originStr or 'logger.debug(' in originStr and ('(' in originStr and ')' in originStr):
        p.insert_paragraph_before('//打印日志', p.style).paragraph_format.left_indent = Pt(20)


def addCommit(p, originStr):
    if 'commit()' in originStr:
        p.insert_paragraph_before('//提交事务 ', p.style).paragraph_format.left_indent = Pt(20)


def addRoolBack(p, originStr):
    if 'rollback()' in originStr:
        p.insert_paragraph_before('//回滚事务 ', p.style).paragraph_format.left_indent = Pt(20)


def addNewClass(p, originStr):
    if '=' in originStr and 'new' in originStr and ' ' in originStr and ';' in originStr:
        nameList = originStr.split(' ')
        for r, n in enumerate(nameList):
            if n == '=':
                className = nameList[r - 1]
                p.insert_paragraph_before('//新建 ' + className + '实例', p.style).paragraph_format.left_indent = Pt(20)


def isFunctionHead(originStr):
    return 'public' in originStr or 'private' in originStr or 'void' in originStr


def isFunctionTail(originStr):
    # 去除{前空格
    handStr = originStr.rstrip()
    if handStr.endswith('{'):
        # print('handStr: '+handStr)
        handedStr = (handStr[:handStr.rindex('{') - 1]).rstrip()
        nameList = handedStr.split(' ')
        for i in nameList:
            if i.startswith('(') and i.endswith(')') or i.endswith('()'):
                return True
    return False


def getNowTime():
    return str(datetime.now().strftime('%Y_%m_%d_%H:%M:%S.%f'))


def getNewClassName():
    pass
